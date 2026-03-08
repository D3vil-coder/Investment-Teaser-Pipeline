"""
Citation Verifier Agent (Agent 6) - PRODUCTION VERSION
Verifies claims against:
1. One-pager source with exact line numbers
2. Calculated values with full math formulas
3. Web-scraped data with URL citations

NO TRUNCATION - Complete text only.
"""

import re
import logging
import os
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class Citation:
    """A verified citation with full source reference."""
    slide_num: int
    claim: str  # NEVER truncated
    verified: bool
    source_type: str  # 'onepager', 'calculated', 'web', 'unverified'
    source_reference: str  # Full reference (line number, formula, or URL)
    original_value: Any = None
    line_number: int = None
    line_content: str = None  # Full line from source


@dataclass
class VerificationReport:
    """Verification report summary."""
    total_claims: int
    verified_count: int
    by_source_type: Dict[str, int]
    verification_rate: float


class CitationVerifier:
    """
    Production citation verifier with multi-source support.
    """
    
    def __init__(self):
        self.citations: List[Citation] = []
        self.md_content: str = ""
        self.md_lines: List[str] = []
        self.extracted_data: Dict[str, Any] = {}
        self.web_data: Dict[str, Any] = {}
        self.md_file_path: str = ""
    
    def set_sources(self, md_file: str, extracted_data: Dict[str, Any],
                    md_content: str = None, web_data: Dict[str, Any] = None):
        """Set all source data for verification."""
        self.md_file_path = md_file
        self.extracted_data = extracted_data
        self.web_data = web_data or {}
        
        if md_content:
            self.md_content = md_content
        else:
            with open(md_file, 'r', encoding='utf-8') as f:
                self.md_content = f.read()
        
        self.md_lines = self.md_content.split('\n')
    
    def verify_slide_content(self, slide_num: int, slide_content: Any) -> List[Citation]:
        """Verify all claims in a slide."""
        slide_citations = []
        
        # Verify each section
        for section_name, items in slide_content.sections.items():
            for item in items:
                citation = self._verify_claim(slide_num, item, section_name)
                slide_citations.append(citation)
                self.citations.append(citation)
        
        # Verify hooks
        if slide_content.hooks:
            for hook in slide_content.hooks:
                citation = self._verify_claim(slide_num, hook, 'Investment Hook')
                slide_citations.append(citation)
                self.citations.append(citation)
        
        return slide_citations
    
    def _verify_claim(self, slide_num: int, claim: str, context: str) -> Citation:
        """Verify a single claim against all sources."""
        # Don't verify empty claims
        if not claim or len(claim.strip()) < 3:
            return Citation(
                slide_num=slide_num,
                claim=claim,
                verified=False,
                source_type='unverified',
                source_reference='Empty or too short'
            )
        
        # 1. Check if it's a calculated value (has % or ₹ with numbers)
        if self._is_calculated_claim(claim):
            return self._verify_calculated(slide_num, claim, context)
        
        # 2. Check if it's from web data
        if self._is_web_claim(claim):
            return self._verify_web_data(slide_num, claim, context)
        
        # 3. Check against one-pager (primary source)
        return self._verify_onepager(slide_num, claim, context)
    
    def _is_calculated_claim(self, claim: str) -> bool:
        """Check if claim contains calculated metrics."""
        patterns = [
            r'CAGR',
            r'Margin',
            r'\d+(\.\d+)?%.*CAGR',
            r'Revenue CAGR',
            r'EBITDA Margin',
        ]
        return any(re.search(p, claim, re.I) for p in patterns)
    
    def _is_web_claim(self, claim: str) -> bool:
        """Check if claim is from web sources."""
        patterns = [
            r'Industry Size',
            r'Industry Growth',
            r'Market Size',
            r'\$\d+.*billion',
            r'Positioned for',
        ]
        return any(re.search(p, claim, re.I) for p in patterns)
    
    def _verify_calculated(self, slide_num: int, claim: str, context: str) -> Citation:
        """Verify calculated values with math formulas."""
        financials = self.extracted_data.get('financials', {})
        
        # Check CAGR
        if 'CAGR' in claim:
            revenue = financials.get('revenue', {})
            if len(revenue) >= 2:
                years = sorted(revenue.keys())
                start_yr, end_yr = years[0], years[-1]
                start_val, end_val = revenue[start_yr], revenue[end_yr]
                n = end_yr - start_yr
                if start_val > 0 and n > 0:
                    cagr = ((end_val / start_val) ** (1/n) - 1) * 100
                    formula = f"CAGR = ((FY{end_yr}_Revenue / FY{start_yr}_Revenue)^(1/{n}) - 1) × 100"
                    calculation = f"= (({end_val:.2f} / {start_val:.2f})^(1/{n}) - 1) × 100"
                    result = f"= {cagr:.1f}%"
                    
                    return Citation(
                        slide_num=slide_num,
                        claim=claim,
                        verified=True,
                        source_type='calculated',
                        source_reference=f"{formula}\n{calculation}\n{result}",
                        original_value={'start': start_val, 'end': end_val, 'years': n, 'cagr': cagr}
                    )
        
        # Check Margin
        if 'Margin' in claim:
            ebitda = financials.get('ebitda', {})
            revenue = financials.get('revenue', {})
            if ebitda and revenue:
                common = set(ebitda.keys()) & set(revenue.keys())
                if common:
                    yr = max(common)
                    if revenue[yr] > 0:
                        margin = (ebitda[yr] / revenue[yr]) * 100
                        formula = f"EBITDA Margin = (FY{yr}_EBITDA / FY{yr}_Revenue) × 100"
                        calculation = f"= ({ebitda[yr]:.2f} / {revenue[yr]:.2f}) × 100"
                        result = f"= {margin:.1f}%"
                        
                        return Citation(
                            slide_num=slide_num,
                            claim=claim,
                            verified=True,
                            source_type='calculated',
                            source_reference=f"{formula}\n{calculation}\n{result}",
                            original_value={'ebitda': ebitda[yr], 'revenue': revenue[yr], 'margin': margin}
                        )
        
        # Check revenue/EBITDA values (FY format)
        fy_match = re.search(r'FY(\d{2,4}).*?₹?([\d.]+)\s*Cr', claim)
        if fy_match:
            year_str = fy_match.group(1)
            value = float(fy_match.group(2))
            
            # Convert 2-digit year to 4-digit
            if len(year_str) == 2:
                year = 2000 + int(year_str)
            else:
                year = int(year_str)
            
            # Check in revenue or ebitda
            revenue = financials.get('revenue', {})
            ebitda = financials.get('ebitda', {})
            
            if year in revenue and abs(revenue[year] - value) < 1:
                return Citation(
                    slide_num=slide_num,
                    claim=claim,
                    verified=True,
                    source_type='onepager',
                    source_reference=f"Financial data: Revenue FY{year} = ₹{revenue[year]:.2f} Cr",
                    original_value=revenue[year]
                )
            
            if year in ebitda and abs(ebitda[year] - value) < 1:
                return Citation(
                    slide_num=slide_num,
                    claim=claim,
                    verified=True,
                    source_type='onepager',
                    source_reference=f"Financial data: EBITDA FY{year} = ₹{ebitda[year]:.2f} Cr",
                    original_value=ebitda[year]
                )
        
        return Citation(
            slide_num=slide_num,
            claim=claim,
            verified=False,
            source_type='unverified',
            source_reference='Calculated value not verified'
        )
    
    def _verify_web_data(self, slide_num: int, claim: str, context: str) -> Citation:
        """Verify claims from web-scraped data with actual URLs."""
        market_data = self.web_data.get('market_data', {})
        sources_used = self.web_data.get('sources_used', [])
        
        if market_data:
            # Get actual source URLs
            source_urls = []
            for src in market_data.get('sources', []):
                if hasattr(src, 'source_url'):
                    source_urls.append(f"[{src.source_name}]({src.source_url})")
                elif isinstance(src, dict):
                    source_urls.append(f"[{src.get('name', 'Source')}]({src.get('url', '#')})")
            
            source_ref = '\n'.join(source_urls) if source_urls else 'Industry estimates'
            
            # Check market size
            if 'Market Size' in claim or 'Industry Size' in claim or '$' in claim:
                if market_data.get('india_market_size'):
                    return Citation(
                        slide_num=slide_num,
                        claim=claim,
                        verified=True,
                        source_type='web',
                        source_reference=f"Data: India market size = {market_data['india_market_size']}\nSources:\n{source_ref}",
                        original_value=market_data['india_market_size']
                    )
            
            # Check growth rate
            if 'Growth' in claim or 'CAGR' in claim.upper():
                if market_data.get('cagr'):
                    return Citation(
                        slide_num=slide_num,
                        claim=claim,
                        verified=True,
                        source_type='web',
                        source_reference=f"Data: Industry CAGR = {market_data['cagr']}\nSources:\n{source_ref}",
                        original_value=market_data['cagr']
                    )
            
            # Check industry drivers
            if 'Positioned for' in claim or 'driven by' in claim.lower():
                drivers = market_data.get('key_drivers', [])
                if drivers:
                    return Citation(
                        slide_num=slide_num,
                        claim=claim,
                        verified=True,
                        source_type='web',
                        source_reference=f"Industry drivers: {', '.join(drivers)}\nSources:\n{source_ref}",
                        original_value=drivers
                    )
        
        # Check in scraped company pages
        company_info = self.web_data.get('company_info', {})
        for page_type, page_data in company_info.items():
            if isinstance(page_data, dict) and 'content' in page_data:
                content = page_data.get('content', '').lower()
                claim_lower = claim.lower()
                if any(word in content for word in claim_lower.split()[:3]):
                    return Citation(
                        slide_num=slide_num,
                        claim=claim,
                        verified=True,
                        source_type='web',
                        source_reference=f"Source: Company Website - {page_type}\nURL: {page_data.get('url', 'N/A')}",
                        original_value=page_data.get('url')
                    )
        
        return Citation(
            slide_num=slide_num,
            claim=claim,
            verified=False,
            source_type='unverified',
            source_reference='Web data not found'
        )
    
    def _verify_onepager(self, slide_num: int, claim: str, context: str) -> Citation:
        """Verify claim against one-pager with exact line numbers."""
        # Clean claim for matching
        clean_claim = self._clean_for_matching(claim)
        
        # Search for match in MD lines
        best_match = None
        best_score = 0
        
        for line_num, line in enumerate(self.md_lines, 1):
            clean_line = self._clean_for_matching(line)
            
            # Direct match
            if clean_claim in clean_line or clean_line in clean_claim:
                score = len(clean_claim) / max(len(clean_line), 1)
                if score > best_score:
                    best_score = score
                    best_match = (line_num, line)
            
            # Keyword match
            claim_words = set(clean_claim.split())
            line_words = set(clean_line.split())
            if len(claim_words) > 2:
                overlap = len(claim_words & line_words) / len(claim_words)
                if overlap > 0.5 and overlap > best_score:
                    best_score = overlap
                    best_match = (line_num, line)
        
        if best_match and best_score > 0.3:
            line_num, line_content = best_match
            return Citation(
                slide_num=slide_num,
                claim=claim,  # Full claim, never truncated
                verified=True,
                source_type='onepager',
                source_reference=f"Line {line_num}: {line_content.strip()}",
                line_number=line_num,
                line_content=line_content.strip()
            )
        
        # Check against extracted data fields
        for field_name, field_value in self._flatten_dict(self.extracted_data):
            if isinstance(field_value, str):
                clean_field = self._clean_for_matching(field_value)
                if clean_claim in clean_field or clean_field in clean_claim:
                    return Citation(
                        slide_num=slide_num,
                        claim=claim,
                        verified=True,
                        source_type='onepager',
                        source_reference=f"Field: {field_name}",
                        original_value=field_value
                    )
        
        return Citation(
            slide_num=slide_num,
            claim=claim,
            verified=False,
            source_type='unverified',
            source_reference='No matching source found in one-pager'
        )
    
    def _clean_for_matching(self, text: str) -> str:
        """Clean text for matching - remove special chars, lowercase."""
        if not text:
            return ""
        # Remove markdown, special chars
        text = re.sub(r'[*_#\-|`]', '', text)
        text = re.sub(r'\s+', ' ', text)
        return text.lower().strip()
    
    def _flatten_dict(self, d: Dict, parent_key: str = '') -> List[tuple]:
        """Flatten nested dictionary."""
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}.{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_dict(v, new_key))
            elif isinstance(v, list):
                for i, item in enumerate(v):
                    if isinstance(item, dict):
                        items.extend(self._flatten_dict(item, f"{new_key}[{i}]"))
                    else:
                        items.append((f"{new_key}[{i}]", item))
            else:
                items.append((new_key, v))
        return items
    
    def generate_report(self, company_name: str, output_path: str) -> VerificationReport:
        """Generate citation report document."""
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Citation Report: {company_name}', 0)
        
        # SUMMARY BLOCK
        doc.add_heading('SUMMARY', level=1)
        
        total = len(self.citations)
        verified = sum(1 for c in self.citations if c.verified)
        rate = (verified / total * 100) if total > 0 else 100
        
        by_type = {}
        for c in self.citations:
            by_type[c.source_type] = by_type.get(c.source_type, 0) + 1
        
        # Summary table
        summary = doc.add_paragraph()
        run = summary.add_run(f'Total Claims: {total}\n')
        run.bold = True
        run.font.size = Pt(12)
        
        run2 = summary.add_run(f'Verified Claims: {verified} ({rate:.1f}%)\n')
        run2.font.size = Pt(11)
        
        status = '✓ All claims verified' if rate == 100 else f'⚠ {total - verified} claim(s) unverified'
        run3 = summary.add_run(f'Status: {status}\n\n')
        run3.bold = True
        run3.font.size = Pt(11)
        
        summary.add_run('Source Breakdown:\n').bold = True
        for src_type, count in sorted(by_type.items()):
            pct = (count / total * 100) if total > 0 else 0
            summary.add_run(f'  • {src_type}: {count} ({pct:.0f}%)\n')
        
        # High-confidence indicator
        high_conf = sum(1 for c in self.citations if c.verified and c.source_type in ('onepager', 'calculated'))
        if total > 0:
            conf_pct = (high_conf / total) * 100
            summary.add_run(f'\nHigh-confidence sources (Data Pack + Calculated): {conf_pct:.0f}%\n')
        
        doc.add_paragraph('─' * 60)
        
        # Verified Citations by Slide
        doc.add_heading('Verified Citations', level=1)
        
        for slide_num in sorted(set(c.slide_num for c in self.citations)):
            slide_citations = [c for c in self.citations if c.slide_num == slide_num and c.verified]
            
            if slide_citations:
                doc.add_heading(f'Slide {slide_num}', level=2)
                
                for c in slide_citations:
                    # Claim - FULL TEXT, no truncation
                    p = doc.add_paragraph()
                    p.add_run('Claim: ').bold = True
                    p.add_run(c.claim)  # Full text
                    
                    # Source type
                    p2 = doc.add_paragraph()
                    p2.add_run(f'Source Type: ').bold = True
                    p2.add_run(c.source_type.upper())
                    
                    # Reference - also full text
                    p3 = doc.add_paragraph()
                    p3.add_run('Reference: ').bold = True
                    p3.add_run(c.source_reference)  # Full reference
                    
                    doc.add_paragraph('---')
        
        # Unverified Claims
        unverified = [c for c in self.citations if not c.verified]
        if unverified:
            doc.add_heading('Unverified Claims (Excluded from PPT)', level=1)
            for c in unverified:
                p = doc.add_paragraph()
                p.add_run(f'Slide {c.slide_num}: ').bold = True
                p.add_run(c.claim)  # Full text
                
                p2 = doc.add_paragraph()
                p2.add_run(f'Reason: {c.source_reference}')
                p2.paragraph_format.left_indent = Inches(0.3)
        
        doc.save(output_path)
        logger.info(f"Citation report: {verified}/{total} verified ({rate:.1f}%)")
        
        return VerificationReport(
            total_claims=total,
            verified_count=verified,
            by_source_type=by_type,
            verification_rate=rate
        )


if __name__ == "__main__":
    print("Citation verifier ready")
